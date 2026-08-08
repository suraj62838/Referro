from django.contrib.auth.models import User
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APITestCase
from rest_framework_simplejwt.tokens import RefreshToken


class AuthTests(APITestCase):

    def test_auth_flow(self):
        """
        Verify signup -> login -> accessing authenticated endpoint.
        Uses /api/health/ (open) and a dummy test endpoint to check auth requirements.
        """
        signup_url = reverse("auth-signup")
        login_url = reverse("auth-token-obtain")
        refresh_url = reverse("auth-token-refresh")

        # 1. Signup a new user
        signup_data = {"email": "testuser@example.com", "password": "securePass123!"}
        response = self.client.post(signup_url, signup_data, format="json")
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn("access", response.data)
        self.assertIn("refresh", response.data)
        self.assertEqual(response.data["user"]["email"], "testuser@example.com")

        # 2. Try to signup with the same email (should fail)
        response_duplicate = self.client.post(
            signup_url, signup_data, format="json"
        )
        self.assertEqual(response_duplicate.status_code, status.HTTP_400_BAD_REQUEST)

        # 3. Login with the credentials
        login_data = {"email": "testuser@example.com", "password": "securePass123!"}
        response = self.client.post(login_url, login_data, format="json")
        if response.status_code != status.HTTP_200_OK:
            print("Login failed with errors:", response.data)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn("access", response.data)
        access_token = response.data["access"]
        refresh_token = response.data["refresh"]

        # 4. Access a protected endpoint without token (should fail with 401)
        # Since we don't have another model endpoint yet, let's call the admin or a dummy
        # view, or we can use the client credentials to test simple jwt.
        # Actually, let's define a simple protected test view for this purpose.
        protected_url = "/api/auth/test-protected/"
        response = self.client.get(protected_url)
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

        # 5. Access the protected endpoint with valid token (should succeed with 200)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {access_token}")
        response = self.client.get(protected_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["message"], "authenticated")

        # 6. Test Refresh Token
        self.client.credentials()  # Clear credentials
        refresh_data = {"refresh": refresh_token}
        response = self.client.post(refresh_url, refresh_data, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn("access", response.data)


class JobCrudTests(APITestCase):

    def setUp(self):
        # Create two users for scoping tests
        self.user1 = User.objects.create_user(username="user1@example.com", email="user1@example.com", password="password123")
        self.user2 = User.objects.create_user(username="user2@example.com", email="user2@example.com", password="password123")

        self.user1.profile.is_verified = True
        self.user1.profile.save()
        self.user2.profile.is_verified = True
        self.user2.profile.save()

        # Generate JWT tokens directly
        self.token1 = str(RefreshToken.for_user(self.user1).access_token)
        self.token2 = str(RefreshToken.for_user(self.user2).access_token)

    def test_job_posting_crud(self):
        postings_url = "/api/job-postings/"

        # 1. Anonymous create should fail
        response = self.client.post(postings_url, {"company_name": "Acme", "role_title": "Engineer"})
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

        # 2. Authenticated user1 creates a posting
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token1}")
        posting_data = {
            "company_name": "Acme",
            "role_title": "Backend Engineer",
            "location": "Remote",
            "recruiter_email": "jobs@acme.com",
            "jd_text": "We need a Django dev",
        }
        response = self.client.post(postings_url, posting_data, format="json")
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        posting_id = response.data["id"]

        # 3. List job postings (public board) - anyone authenticated can view
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token2}")
        response = self.client.get(postings_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data), 1)

        # 4. User2 tries to update user1's posting (should fail)
        detail_url = f"{postings_url}{posting_id}/"
        response = self.client.patch(detail_url, {"location": "New York"}, format="json")
        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

        # 5. User1 updates their own posting (should succeed)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token1}")
        response = self.client.patch(detail_url, {"location": "New York"}, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["location"], "New York")

        # 6. User2 tries to delete user1's posting (should fail)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token2}")
        response = self.client.delete(detail_url)
        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

        # 7. User1 deletes their own posting (should succeed)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token1}")
        response = self.client.delete(detail_url)
        self.assertEqual(response.status_code, status.HTTP_204_NO_CONTENT)

    def test_job_application_crud(self):
        apps_url = "/api/job-applications/"

        # 1. Create application for user1
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token1}")
        app_data = {
            "company_name": "Google",
            "role_title": "Frontend dev",
            "recruiter_email": "recruiter@google.com",
            "jd_text": "React experience required",
        }
        response = self.client.post(apps_url, app_data, format="json")
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        app_id = response.data["id"]

        # 2. User2 lists applications (should not see User1's application)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token2}")
        response = self.client.get(apps_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data), 0)

        # 3. User2 tries to retrieve User1's application detail (should fail with 404)
        detail_url = f"{apps_url}{app_id}/"
        response = self.client.get(detail_url)
        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)

        # 4. User1 retrieves detail (succeeds)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token1}")
        response = self.client.get(detail_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["company_name"], "Google")

        # 5. User1 updates status
        response = self.client.patch(detail_url, {"status": "interview"}, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["status"], "interview")

        # 6. User1 deletes application
        response = self.client.delete(detail_url)
        self.assertEqual(response.status_code, status.HTTP_204_NO_CONTENT)


class ExtractTests(APITestCase):
    """Phase 3: Tests for POST /api/job-applications/extract/."""

    EXTRACT_URL = "/api/job-applications/extract/"

    def setUp(self):
        self.user = User.objects.create_user(
            username="extractor@example.com",
            email="extractor@example.com",
            password="securePass123!",
        )
        self.user.profile.is_verified = True
        self.user.profile.save()
        self.token = str(RefreshToken.for_user(self.user).access_token)

    def _auth(self):
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token}")

    # ── Text extraction tests ─────────────────────────────────

    def test_text_with_email(self):
        """Text containing an email returns both jd_text and recruiter_email."""
        self._auth()
        jd = "We are hiring! Contact us at jobs@acme.com for details."
        res = self.client.post(self.EXTRACT_URL, {"text": jd}, format="json")
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertEqual(res.data["jd_text"], jd)
        self.assertEqual(res.data["recruiter_email"], "jobs@acme.com")

    def test_text_without_email(self):
        """Text without an email returns jd_text with empty recruiter_email."""
        self._auth()
        jd = "Join our team — no contact info here."
        res = self.client.post(self.EXTRACT_URL, {"text": jd}, format="json")
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertEqual(res.data["jd_text"], jd)
        self.assertEqual(res.data["recruiter_email"], "")

    def test_no_input(self):
        """Neither text nor file → 400."""
        self._auth()
        res = self.client.post(self.EXTRACT_URL, {}, format="json")
        self.assertEqual(res.status_code, status.HTTP_400_BAD_REQUEST)

    def test_unauthenticated(self):
        """Unauthenticated request → 401."""
        res = self.client.post(
            self.EXTRACT_URL, {"text": "hello"}, format="json"
        )
        self.assertEqual(res.status_code, status.HTTP_401_UNAUTHORIZED)

    # ── File extraction tests ─────────────────────────────────

    def test_pdf_upload(self):
        """PDF file upload → returns extracted text + email."""
        import io

        from PyPDF2 import PdfWriter
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        # Create a real PDF with text
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.drawString(72, 720, "Backend Engineer at Acme Corp")
        c.drawString(72, 700, "Contact: hiring@acmecorp.com")
        c.save()
        buf.seek(0)

        from django.core.files.uploadedfile import SimpleUploadedFile

        pdf_file = SimpleUploadedFile(
            "job.pdf", buf.read(), content_type="application/pdf"
        )

        self._auth()
        res = self.client.post(
            self.EXTRACT_URL, {"file": pdf_file}, format="multipart"
        )
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertIn("jd_text", res.data)
        self.assertTrue(len(res.data["jd_text"]) > 0)
        # The email should be detected from the extracted text
        self.assertEqual(res.data["recruiter_email"], "hiring@acmecorp.com")

    def test_docx_upload(self):
        """DOCX file upload → returns extracted text + email."""
        import io

        import docx

        doc = docx.Document()
        doc.add_paragraph("Frontend Developer at Widget Inc.")
        doc.add_paragraph("Apply at apply@widgetinc.com")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        from django.core.files.uploadedfile import SimpleUploadedFile

        docx_file = SimpleUploadedFile(
            "job.docx",
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        self._auth()
        res = self.client.post(
            self.EXTRACT_URL, {"file": docx_file}, format="multipart"
        )
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertIn("Widget", res.data["jd_text"])
        self.assertEqual(res.data["recruiter_email"], "apply@widgetinc.com")

    def test_unsupported_file_type(self):
        """Uploading an unsupported file type → 400."""
        from django.core.files.uploadedfile import SimpleUploadedFile

        txt_file = SimpleUploadedFile(
            "notes.txt", b"just a text file", content_type="text/plain"
        )
        self._auth()
        res = self.client.post(
            self.EXTRACT_URL, {"file": txt_file}, format="multipart"
        )
        self.assertEqual(res.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("Unsupported", res.data["detail"])


class Phase4AITests(APITestCase):
    """Phase 4: Tests for Job Description generation and email drafting."""

    def setUp(self):
        self.user = User.objects.create_user(
            username="ai_user@example.com",
            email="ai_user@example.com",
            password="securePass123!",
        )
        self.user.profile.is_verified = True
        self.user.profile.save()
        self.token = str(RefreshToken.for_user(self.user).access_token)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token}")

    def test_generate_jd_endpoint(self):
        """POST /api/job-postings/generate-jd/ creates a JD with AI service."""
        url = "/api/job-postings/generate-jd/"
        data = {
            "role_title": "React Developer",
            "seniority": "Senior",
            "key_skills": "React, TypeScript, CSS",
            "notes": "Remote, small team",
        }
        res = self.client.post(url, data, format="json")
        if res.status_code != status.HTTP_200_OK:
            print("Generate JD response:", res.data)
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertIn("jd_text", res.data)
        self.assertIn("React Developer", res.data["jd_text"])
        self.assertIn("Senior", res.data["jd_text"])
        self.assertIn("React, TypeScript, CSS", res.data["jd_text"])

    def test_generate_jd_missing_role(self):
        """POST /api/job-postings/generate-jd/ fails if role is missing."""
        url = "/api/job-postings/generate-jd/"
        data = {
            "seniority": "Senior",
            "key_skills": "React",
        }
        res = self.client.post(url, data, format="json")
        self.assertEqual(res.status_code, status.HTTP_400_BAD_REQUEST)

    def test_draft_email_endpoint(self):
        """POST /api/job-applications/{id}/draft-email/ drafts email for application."""
        from api.models import JobApplication
        app = JobApplication.objects.create(
            user=self.user,
            company_name="Stripe",
            role_title="Backend Dev",
            jd_text="Python knowledge required.",
            recruiter_email="jobs@stripe.com",
        )
        url = f"/api/job-applications/{app.id}/draft-email/"
        res = self.client.post(url, format="json")
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertIn("subject", res.data)
        self.assertIn("body", res.data)
        self.assertIn("Stripe", res.data["subject"])
        self.assertIn("Backend Dev", res.data["subject"])


class Phase5OAuthAndSendingTests(APITestCase):
    """Phase 5: Tests for Google OAuth flows, token storage, and Gmail sending."""

    def setUp(self):
        self.user = User.objects.create_user(
            username="oauth_user@example.com",
            email="oauth_user@example.com",
            password="securePass123!",
        )
        self.user.profile.is_verified = True
        self.user.profile.save()
        self.token = str(RefreshToken.for_user(self.user).access_token)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token}")

    def test_oauth_connect_endpoint(self):
        """GET /api/email-accounts/oauth/connect/ returns authorization redirect URL."""
        url = "/api/email-accounts/oauth/connect/"
        res = self.client.get(url)
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertIn("auth_url", res.data)
        self.assertIn("accounts.google.com", res.data["auth_url"])
        self.assertIn("state=", res.data["auth_url"])

    def test_email_account_me_endpoint_not_connected(self):
        """GET /api/email-accounts/me/ returns connected: False if no account exists."""
        url = "/api/email-accounts/me/"
        res = self.client.get(url)
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertFalse(res.data["connected"])
        self.assertIsNone(res.data["email"])

    def test_oauth_callback_endpoint(self):
        """GET /api/email-accounts/oauth/callback/ handles state/code, stores tokens, redirects."""
        from unittest.mock import patch
        from rest_framework_simplejwt.tokens import AccessToken

        # Generate state JWT
        token = AccessToken()
        token["user_id"] = self.user.id
        state = str(token)

        callback_url = "/api/email-accounts/oauth/callback/"
        
        with patch("services.gmail_service.exchange_code") as mock_exchange:
            mock_exchange.return_value = {
                "access_token": "mock_access_token_123",
                "refresh_token": "mock_refresh_token_456",
                "email": "connected_gmail@example.com",
            }

            res = self.client.get(f"{callback_url}?code=google_code_xyz&state={state}")
            
            # Verify it redirects to the frontend dashboard with connection confirmation
            self.assertEqual(res.status_code, 302)
            self.assertIn("/dashboard?connected=1", res.url)

            # Verify EmailAccount was created
            from api.models import EmailAccount
            account = EmailAccount.objects.filter(user=self.user).first()
            self.assertIsNotNone(account)
            self.assertEqual(account.email_address, "connected_gmail@example.com")
            # Verify property getter/setter work through encryption
            self.assertEqual(account.access_token, "mock_access_token_123")
            self.assertEqual(account.refresh_token, "mock_refresh_token_456")

            # Check email_account_me now shows connected
            me_res = self.client.get("/api/email-accounts/me/")
            self.assertEqual(me_res.status_code, status.HTTP_200_OK)
            self.assertTrue(me_res.data["connected"])
            self.assertEqual(me_res.data["email"], "connected_gmail@example.com")

    def test_send_email_endpoint(self):
        """POST /api/job-applications/{id}/send-email/ sends outreach via Gmail and logs it."""
        from unittest.mock import patch
        from api.models import EmailAccount, JobApplication, EmailLog

        # 1. Create a job application
        app = JobApplication.objects.create(
            user=self.user,
            company_name="Google",
            role_title="Backend Engineer",
            recruiter_email="recruiter@google.com",
            jd_text="Job details...",
            status="sent", # defaults to sent, but we will send
        )

        send_url = f"/api/job-applications/{app.id}/send-email/"

        # Try to send without a connected email account first
        res_no_account = self.client.post(
            send_url,
            {"subject": "Hello", "body": "My pitch"},
            format="json"
        )
        self.assertEqual(res_no_account.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("connect your Gmail", res_no_account.data["detail"])

        # Create connected account
        account = EmailAccount.objects.create(
            user=self.user,
            provider="gmail",
            email_address="sender@gmail.com",
        )
        account.access_token = "access"
        account.refresh_token = "refresh"
        account.save()

        with patch("services.gmail_service.send_email") as mock_send:
            mock_send.return_value = "thread_id_abc_123"

            res = self.client.post(
                send_url,
                {"subject": "Hello recruiter", "body": "I am interested"},
                format="json"
            )
            self.assertEqual(res.status_code, status.HTTP_200_OK)
            self.assertTrue(res.data["success"])
            self.assertEqual(res.data["thread_id"], "thread_id_abc_123")

            # Verify EmailLog was created
            log = EmailLog.objects.filter(job_application=app).first()
            self.assertIsNotNone(log)
            self.assertEqual(log.subject, "Hello recruiter")
            self.assertEqual(log.body, "I am interested")
            self.assertEqual(log.gmail_thread_id, "thread_id_abc_123")

            # Verify application status was updated/remains sent (or actually set to sent)
            app.refresh_from_db()
            self.assertEqual(app.status, "sent")


class Phase6ReplyPollingTests(APITestCase):
    """Phase 6: Tests for reply polling task and detail view nested logs."""

    def setUp(self):
        self.user = User.objects.create_user(
            username="poll_user@example.com",
            email="poll_user@example.com",
            password="securePass123!",
        )
        self.user.profile.is_verified = True
        self.user.profile.save()
        self.token = str(RefreshToken.for_user(self.user).access_token)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token}")

    def test_application_detail_nested_logs(self):
        """GET /api/job-applications/{id}/ includes email_logs and reply_logs."""
        from api.models import EmailLog, JobApplication, ReplyLog

        app = JobApplication.objects.create(
            user=self.user,
            company_name="Netflix",
            role_title="Senior Engineer",
            status="sent",
        )
        email_log = EmailLog.objects.create(
            job_application=app,
            subject="Application for Senior Engineer",
            body="Hello, I would like to apply.",
            gmail_thread_id="thread_xyz_789",
        )
        reply_log = ReplyLog.objects.create(
            job_application=app,
            snippet="Thanks for reaching out!",
            body="Thanks for reaching out! We'd love to chat.",
            gmail_message_id="msg_reply_123",
        )

        res = self.client.get(f"/api/job-applications/{app.id}/")
        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertIn("email_logs", res.data)
        self.assertIn("reply_logs", res.data)
        self.assertEqual(len(res.data["email_logs"]), 1)
        self.assertEqual(len(res.data["reply_logs"]), 1)
        self.assertEqual(res.data["email_logs"][0]["subject"], "Application for Senior Engineer")
        self.assertEqual(res.data["reply_logs"][0]["snippet"], "Thanks for reaching out!")

    def test_poll_replies_task(self):
        """poll_replies task detects new messages in sent thread, creates ReplyLog, updates status to replied."""
        from unittest.mock import patch
        from api.models import EmailAccount, EmailLog, JobApplication, ReplyLog
        from api.tasks import poll_replies

        # 1. Setup email account and sent application
        account = EmailAccount.objects.create(
            user=self.user,
            provider="gmail",
            email_address="applicant@example.com",
        )
        account.access_token = "valid_token"
        account.save()

        app = JobApplication.objects.create(
            user=self.user,
            company_name="Apple",
            role_title="iOS Developer",
            status="sent",
        )

        email_log = EmailLog.objects.create(
            job_application=app,
            subject="Outreach for iOS Dev",
            body="Hi, sending my resume.",
            gmail_thread_id="apple_thread_1",
        )

        # Mock get_thread_messages to return two messages: 1 sent by user, 1 received reply
        mock_messages = [
            {
                "id": "msg_sent_1",
                "snippet": "Hi, sending my resume.",
                "body": "Hi, sending my resume.",
                "date": "Mon, 27 Jul 2026 10:00:00 GMT",
                "from": "applicant@example.com",
            },
            {
                "id": "msg_reply_2",
                "snippet": "We saw your resume and would like an interview.",
                "body": "We saw your resume and would like an interview.",
                "date": "Mon, 27 Jul 2026 12:00:00 GMT",
                "from": "recruiter@apple.com",
            },
        ]

        with patch("services.gmail_service.get_thread_messages") as mock_get_thread:
            mock_get_thread.return_value = mock_messages

            result = poll_replies()
            self.assertIn("found 1 new replies", result)

            # Check that ReplyLog was created
            replies = ReplyLog.objects.filter(job_application=app)
            self.assertEqual(replies.count(), 1)
            self.assertEqual(replies.first().gmail_message_id, "msg_reply_2")
            self.assertEqual(replies.first().snippet, "We saw your resume and would like an interview.")

            # Check that JobApplication status updated to replied
            app.refresh_from_db()
            self.assertEqual(app.status, "replied")

            # Run poll_replies a second time — application status is now 'replied', so it is no longer polled
            result2 = poll_replies()
            self.assertEqual(result2, "No threads to check.")
            self.assertEqual(ReplyLog.objects.filter(job_application=app).count(), 1)

    def test_check_replies_endpoint(self):
        """The on-demand endpoint records replies without Celery Beat."""
        from unittest.mock import patch
        from api.models import EmailAccount, EmailLog, JobApplication, ReplyLog

        EmailAccount.objects.create(
            user=self.user,
            email_address="poll_user@example.com",
            provider="gmail",
            access_token="token",
            refresh_token="refresh",
        )
        app = JobApplication.objects.create(
            user=self.user,
            company_name="Acme",
            role_title="Engineer",
            status="sent",
        )
        EmailLog.objects.create(job_application=app, subject="Hello", body="Hi", gmail_thread_id="thread_1")

        messages = [
            {"id": "sent_1", "from": "user1@example.com"},
            {"id": "reply_1", "snippet": "Interested", "body": "Let's talk.", "from": "hr@acme.com"},
        ]
        with patch("services.gmail_service.get_thread_messages", return_value=messages):
            res = self.client.post(f"/api/job-applications/{app.id}/check-replies/")

        self.assertEqual(res.status_code, status.HTTP_200_OK)
        self.assertEqual(res.data["new_replies"], 1)
        self.assertEqual(ReplyLog.objects.filter(job_application=app).count(), 1)


class Phase7IntegrationTests(APITestCase):
    """Phase 7: End-to-end flow integration tests, security audit, and input validation."""

    def setUp(self):
        self.user1 = User.objects.create_user(
            username="applicant@example.com",
            email="applicant@example.com",
            password="securePass123!",
        )
        self.user2 = User.objects.create_user(
            username="poster@example.com",
            email="poster@example.com",
            password="securePass123!",
        )
        self.user1.profile.is_verified = True
        self.user1.profile.save()
        self.user2.profile.is_verified = True
        self.user2.profile.save()

        self.token1 = str(RefreshToken.for_user(self.user1).access_token)
        self.token2 = str(RefreshToken.for_user(self.user2).access_token)

    def _auth(self, token):
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    def test_e2e_applicant_flow(self):
        """README §5: Signup -> Extract -> Create -> Draft -> Send -> Poll -> Detail."""
        from unittest.mock import patch
        from api.models import EmailAccount, EmailLog, JobApplication, ReplyLog
        from api.tasks import poll_replies

        self._auth(self.token1)

        # 1. Extract text and email from paste
        jd_text_input = "We are seeking a Python Engineer at Acme Corp. Contact recruiter at hr@acmecorp.com for details."
        res_extract = self.client.post("/api/job-applications/extract/", {"text": jd_text_input}, format="json")
        self.assertEqual(res_extract.status_code, status.HTTP_200_OK)
        self.assertEqual(res_extract.data["recruiter_email"], "hr@acmecorp.com")

        # 2. Create JobApplication
        app_data = {
            "company_name": "Acme Corp",
            "role_title": "Python Engineer",
            "jd_text": res_extract.data["jd_text"],
            "recruiter_email": res_extract.data["recruiter_email"],
        }
        res_create = self.client.post("/api/job-applications/", app_data, format="json")
        self.assertEqual(res_create.status_code, status.HTTP_201_CREATED)
        app_id = res_create.data["id"]

        # 3. Draft email with AI
        res_draft = self.client.post(f"/api/job-applications/{app_id}/draft-email/", format="json")
        self.assertEqual(res_draft.status_code, status.HTTP_200_OK)
        subject = res_draft.data["subject"]
        body = res_draft.data["body"]

        # 4. Connect Gmail account
        account = EmailAccount.objects.create(
            user=self.user1,
            provider="gmail",
            email_address="applicant@example.com",
        )
        account.access_token = "access"
        account.refresh_token = "refresh"
        account.save()

        # 5. Send outreach email
        with patch("services.gmail_service.send_email") as mock_send:
            mock_send.return_value = "thread_e2e_100"
            res_send = self.client.post(
                f"/api/job-applications/{app_id}/send-email/",
                {"subject": subject, "body": body},
                format="json",
            )
            self.assertEqual(res_send.status_code, status.HTTP_200_OK)
            self.assertEqual(res_send.data["thread_id"], "thread_e2e_100")

        # Verify status is sent
        app = JobApplication.objects.get(id=app_id)
        self.assertEqual(app.status, "sent")

        # 6. Poll replies via Celery task
        mock_messages = [
            {"id": "sent_1", "snippet": "Sent email", "body": body, "from": "applicant@example.com"},
            {"id": "reply_1", "snippet": "Let us talk!", "body": "Let us talk on Tuesday.", "from": "hr@acmecorp.com"},
        ]
        with patch("services.gmail_service.get_thread_messages") as mock_get_thread:
            mock_get_thread.return_value = mock_messages
            poll_result = poll_replies()
            self.assertIn("found 1 new replies", poll_result)

        # 7. Check Detail endpoint
        res_detail = self.client.get(f"/api/job-applications/{app_id}/")
        self.assertEqual(res_detail.status_code, status.HTTP_200_OK)
        self.assertEqual(res_detail.data["status"], "replied")
        self.assertEqual(len(res_detail.data["email_logs"]), 1)
        self.assertEqual(len(res_detail.data["reply_logs"]), 1)
        self.assertEqual(res_detail.data["reply_logs"][0]["snippet"], "Let us talk!")

    def test_e2e_job_posting_flow_and_self_application(self):
        """README §6: AI-generate JD -> Create Posting -> List -> Self-apply warning."""
        self._auth(self.token2)

        # 1. AI-generate JD
        gen_data = {
            "role_title": "Fullstack Dev",
            "seniority": "Lead",
            "key_skills": "Django, React",
        }
        res_gen = self.client.post("/api/job-postings/generate-jd/", gen_data, format="json")
        self.assertEqual(res_gen.status_code, status.HTTP_200_OK)
        jd_text = res_gen.data["jd_text"]

        # 2. Create JobPosting
        posting_data = {
            "company_name": "Poster Co",
            "role_title": "Fullstack Dev",
            "jd_text": jd_text,
            "recruiter_email": "poster@example.com",
            "location": "Remote",
        }
        res_post = self.client.post("/api/job-postings/", posting_data, format="json")
        self.assertEqual(res_post.status_code, status.HTTP_201_CREATED)
        posting_id = res_post.data["id"]

        # 3. List postings on board
        res_list = self.client.get("/api/job-postings/")
        self.assertEqual(res_list.status_code, status.HTTP_200_OK)
        self.assertTrue(any(p["id"] == posting_id for p in res_list.data))

        # 4. Poster applies to their own posting (Self-application warning check)
        self_app_data = {
            "job_posting": posting_id,
            "company_name": "Poster Co",
            "role_title": "Fullstack Dev",
            "jd_text": jd_text,
            "recruiter_email": "poster@example.com",
        }
        res_self_app = self.client.post("/api/job-applications/", self_app_data, format="json")
        self.assertEqual(res_self_app.status_code, status.HTTP_201_CREATED)
        self.assertTrue(res_self_app.data["is_self_application"])
        self.assertIsNotNone(res_self_app.data["warning"])

    def test_permission_scoping(self):
        """Cross-user permission isolation audit."""
        from api.models import JobApplication, JobPosting

        # Create posting and application by user2
        posting = JobPosting.objects.create(posted_by=self.user2, company_name="User2 Co", role_title="Dev")
        app = JobApplication.objects.create(user=self.user2, company_name="User2 Co", role_title="Dev")

        # User1 attempts to retrieve User2's application -> 404 (scoped queryset)
        self._auth(self.token1)
        res_app_get = self.client.get(f"/api/job-applications/{app.id}/")
        self.assertEqual(res_app_get.status_code, status.HTTP_404_NOT_FOUND)

        # User1 attempts to update User2's application -> 404
        res_app_patch = self.client.patch(f"/api/job-applications/{app.id}/", {"status": "interview"}, format="json")
        self.assertEqual(res_app_patch.status_code, status.HTTP_404_NOT_FOUND)

        # User1 attempts to update User2's posting -> 403 Forbidden
        res_post_patch = self.client.patch(f"/api/job-postings/{posting.id}/", {"location": "Hacked"}, format="json")
        self.assertEqual(res_post_patch.status_code, status.HTTP_403_FORBIDDEN)

        # User1 attempts to delete User2's posting -> 403 Forbidden
        res_post_delete = self.client.delete(f"/api/job-postings/{posting.id}/")
        self.assertEqual(res_post_delete.status_code, status.HTTP_403_FORBIDDEN)

    def test_input_validation(self):
        """File size limits and email format validation tests."""
        from django.core.files.uploadedfile import SimpleUploadedFile

        self._auth(self.token1)

        # 1. Invalid recruiter email format -> 400
        res_bad_email = self.client.post("/api/job-postings/", {
            "company_name": "Test",
            "role_title": "Dev",
            "recruiter_email": "not-an-email",
        }, format="json")
        self.assertEqual(res_bad_email.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("recruiter_email", res_bad_email.data)

        # 2. File size > 5MB -> 400
        large_content = b"a" * (6 * 1024 * 1024)  # 6MB
        large_file = SimpleUploadedFile("big.pdf", large_content, content_type="application/pdf")
        res_big_file = self.client.post("/api/job-applications/extract/", {"file": large_file}, format="multipart")
        self.assertEqual(res_big_file.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("File too large", res_big_file.data["detail"])


class Phase8EmailVerificationTests(APITestCase):
    """Phase 8: Tests for email verification codes, permissions, resend, and flow."""

    def test_signup_creates_unverified_user_and_code(self):
        """Signup creates user with is_verified=False and creates EmailVerificationCode."""
        from api.models import EmailVerificationCode, UserProfile

        signup_url = reverse("auth-signup")
        data = {"email": "newuser@example.com", "password": "securePass123!"}
        res = self.client.post(signup_url, data, format="json")
        self.assertEqual(res.status_code, status.HTTP_201_CREATED)
        self.assertFalse(res.data["user"]["is_verified"])

        user = User.objects.get(email="newuser@example.com")
        self.assertFalse(user.profile.is_verified)

        code_obj = EmailVerificationCode.objects.filter(user=user).first()
        self.assertIsNotNone(code_obj)
        self.assertEqual(len(code_obj.code), 6)
        self.assertTrue(code_obj.code.isdigit())
        self.assertFalse(code_obj.is_used)

    def test_unverified_user_blocked_from_protected_endpoint(self):
        """Unverified user receives 403 on protected endpoint."""
        user = User.objects.create_user(
            username="unverified@example.com",
            email="unverified@example.com",
            password="securePass123!",
        )
        token = str(RefreshToken.for_user(user).access_token)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        res = self.client.get("/api/job-postings/")
        self.assertEqual(res.status_code, status.HTTP_403_FORBIDDEN)
        self.assertIn("Email not verified", res.data["detail"])

    def test_verify_with_correct_code(self):
        """POST /api/auth/verify-email/ with correct code verifies user and unblocks endpoints."""
        from api.models import EmailVerificationCode

        signup_url = reverse("auth-signup")
        res_signup = self.client.post(
            signup_url,
            {"email": "verifier@example.com", "password": "securePass123!"},
            format="json",
        )
        token = res_signup.data["access"]
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        user = User.objects.get(email="verifier@example.com")
        code_obj = EmailVerificationCode.objects.filter(user=user).first()

        verify_url = reverse("auth-verify-email")
        res_verify = self.client.post(verify_url, {"code": code_obj.code}, format="json")
        self.assertEqual(res_verify.status_code, status.HTTP_200_OK)
        self.assertTrue(res_verify.data["is_verified"])

        user.refresh_from_db()
        self.assertTrue(user.profile.is_verified)

        # Now protected endpoint works
        res_postings = self.client.get("/api/job-postings/")
        self.assertEqual(res_postings.status_code, status.HTTP_200_OK)

    def test_verify_with_wrong_code(self):
        """Wrong code returns 400 error."""
        signup_url = reverse("auth-signup")
        res_signup = self.client.post(
            signup_url,
            {"email": "wrongcode@example.com", "password": "securePass123!"},
            format="json",
        )
        token = res_signup.data["access"]
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        verify_url = reverse("auth-verify-email")
        res_verify = self.client.post(verify_url, {"code": "000000"}, format="json")
        self.assertEqual(res_verify.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("Incorrect verification code", res_verify.data["detail"])

    def test_verify_with_expired_code(self):
        """Expired code returns 400 error."""
        from datetime import timedelta
        from django.utils import timezone
        from api.models import EmailVerificationCode

        signup_url = reverse("auth-signup")
        res_signup = self.client.post(
            signup_url,
            {"email": "expired@example.com", "password": "securePass123!"},
            format="json",
        )
        token = res_signup.data["access"]
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        user = User.objects.get(email="expired@example.com")
        code_obj = EmailVerificationCode.objects.filter(user=user).first()
        code_obj.expires_at = timezone.now() - timedelta(minutes=1)
        code_obj.save()

        verify_url = reverse("auth-verify-email")
        res_verify = self.client.post(verify_url, {"code": code_obj.code}, format="json")
        self.assertEqual(res_verify.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("No valid verification code found", res_verify.data["detail"])

    def test_verify_with_used_code(self):
        """Re-using a code returns 400 error."""
        from api.models import EmailVerificationCode

        signup_url = reverse("auth-signup")
        res_signup = self.client.post(
            signup_url,
            {"email": "reused@example.com", "password": "securePass123!"},
            format="json",
        )
        token = res_signup.data["access"]
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        user = User.objects.get(email="reused@example.com")
        code_obj = EmailVerificationCode.objects.filter(user=user).first()

        verify_url = reverse("auth-verify-email")
        # First use
        res1 = self.client.post(verify_url, {"code": code_obj.code}, format="json")
        self.assertEqual(res1.status_code, status.HTTP_200_OK)

        # Un-verify user profile to test reusing the code
        user.profile.is_verified = False
        user.profile.save()

        # Second use attempt
        res2 = self.client.post(verify_url, {"code": code_obj.code}, format="json")
        self.assertEqual(res2.status_code, status.HTTP_400_BAD_REQUEST)

    def test_resend_code_invalidates_old_codes(self):
        """Resend code invalidates prior unused codes and generates new code."""
        from datetime import timedelta
        from django.utils import timezone
        from api.models import EmailVerificationCode

        signup_url = reverse("auth-signup")
        res_signup = self.client.post(
            signup_url,
            {"email": "resender@example.com", "password": "securePass123!"},
            format="json",
        )
        token = res_signup.data["access"]
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        user = User.objects.get(email="resender@example.com")
        old_code_obj = EmailVerificationCode.objects.filter(user=user).first()

        # Simulate time passing >60s so cooldown allows resend
        old_code_obj.created_at = timezone.now() - timedelta(seconds=65)
        old_code_obj.save()

        resend_url = reverse("auth-resend-code")
        res_resend = self.client.post(resend_url, format="json")
        self.assertEqual(res_resend.status_code, status.HTTP_200_OK)

        old_code_obj.refresh_from_db()
        self.assertTrue(old_code_obj.is_used)

        new_code_obj = (
            EmailVerificationCode.objects.filter(user=user, is_used=False)
            .order_by("-created_at")
            .first()
        )
        self.assertIsNotNone(new_code_obj)
        self.assertNotEqual(new_code_obj.code, old_code_obj.code)

    def test_resend_code_cooldown(self):
        """Resending within 60 seconds returns 429 rate limit."""
        signup_url = reverse("auth-signup")
        res_signup = self.client.post(
            signup_url,
            {"email": "spammer@example.com", "password": "securePass123!"},
            format="json",
        )
        token = res_signup.data["access"]
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

        resend_url = reverse("auth-resend-code")
        res = self.client.post(resend_url, format="json")
        self.assertEqual(res.status_code, status.HTTP_429_TOO_MANY_REQUESTS)
        self.assertIn("Please wait", res.data["detail"])


class Phase9ResumeTests(APITestCase):
    """Phase 9: Tests for resume upload, replace, delete, and email auto-attachment."""

    def setUp(self):
        self.user = User.objects.create_user(
            username="resume_user@example.com",
            email="resume_user@example.com",
            password="securePass123!",
        )
        self.user.profile.is_verified = True
        self.user.profile.save()
        self.token = str(RefreshToken.for_user(self.user).access_token)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {self.token}")

    def test_get_resume_404_when_none_uploaded(self):
        """GET /api/resume/ returns 404 if no resume has been uploaded."""
        res = self.client.get("/api/resume/")
        self.assertEqual(res.status_code, status.HTTP_404_NOT_FOUND)

    def test_upload_and_get_resume(self):
        """POST /api/resume/ uploads PDF, GET /api/resume/ returns metadata."""
        from django.core.files.uploadedfile import SimpleUploadedFile

        dummy_pdf = SimpleUploadedFile(
            "my_resume.pdf",
            b"%PDF-1.4 dummy content",
            content_type="application/pdf",
        )
        res = self.client.post("/api/resume/", {"file": dummy_pdf}, format="multipart")
        self.assertEqual(res.status_code, status.HTTP_201_CREATED)
        self.assertEqual(res.data["original_filename"], "my_resume.pdf")

        # GET should return the metadata
        get_res = self.client.get("/api/resume/")
        self.assertEqual(get_res.status_code, status.HTTP_200_OK)
        self.assertEqual(get_res.data["original_filename"], "my_resume.pdf")

    def test_upload_invalid_file_type(self):
        """Uploading an unallowed file type returns 400."""
        from django.core.files.uploadedfile import SimpleUploadedFile

        txt_file = SimpleUploadedFile(
            "resume.txt", b"plain text", content_type="text/plain"
        )
        res = self.client.post("/api/resume/", {"file": txt_file}, format="multipart")
        self.assertEqual(res.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("Unsupported file type", res.data["detail"])

    def test_upload_exceeds_size_limit(self):
        """Uploading a file > 5MB returns 400."""
        from django.core.files.uploadedfile import SimpleUploadedFile

        large_bytes = b"0" * (5 * 1024 * 1024 + 1)
        large_file = SimpleUploadedFile(
            "big_resume.pdf", large_bytes, content_type="application/pdf"
        )
        res = self.client.post("/api/resume/", {"file": large_file}, format="multipart")
        self.assertEqual(res.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("File too large", res.data["detail"])

    def test_replace_resume_deletes_old_file(self):
        """Uploading a second resume replaces the old one and deletes the old file."""
        from django.core.files.uploadedfile import SimpleUploadedFile
        from api.models import Resume

        pdf1 = SimpleUploadedFile("resume_v1.pdf", b"%PDF-1.4 v1", content_type="application/pdf")
        self.client.post("/api/resume/", {"file": pdf1}, format="multipart")

        r1 = Resume.objects.filter(user=self.user).first()
        self.assertIsNotNone(r1)
        old_file_path = r1.file.path

        # Upload second resume
        pdf2 = SimpleUploadedFile("resume_v2.pdf", b"%PDF-1.4 v2", content_type="application/pdf")
        res2 = self.client.post("/api/resume/", {"file": pdf2}, format="multipart")
        self.assertEqual(res2.status_code, status.HTTP_201_CREATED)

        # Confirm DB has only 1 resume for user
        resumes = Resume.objects.filter(user=self.user)
        self.assertEqual(resumes.count(), 1)
        self.assertEqual(resumes.first().original_filename, "resume_v2.pdf")

        # Confirm old physical file was purged from disk
        import os
        self.assertFalse(os.path.exists(old_file_path))

    def test_delete_resume(self):
        """DELETE /api/resume/ removes the resume record and file."""
        from django.core.files.uploadedfile import SimpleUploadedFile
        from api.models import Resume
        import os

        pdf = SimpleUploadedFile("delete_me.pdf", b"%PDF-1.4 to delete", content_type="application/pdf")
        self.client.post("/api/resume/", {"file": pdf}, format="multipart")

        file_path = Resume.objects.filter(user=self.user).first().file.path

        res_del = self.client.delete("/api/resume/")
        self.assertEqual(res_del.status_code, status.HTTP_200_OK)

        self.assertEqual(Resume.objects.filter(user=self.user).count(), 0)
        self.assertFalse(os.path.exists(file_path))

    def test_send_email_attaches_active_resume(self):
        """send-email automatically attaches the active resume and links it on EmailLog."""
        from unittest.mock import patch
        from django.core.files.uploadedfile import SimpleUploadedFile
        from api.models import EmailAccount, EmailLog, JobApplication, Resume

        # 1. Upload resume
        pdf = SimpleUploadedFile("attached_resume.pdf", b"%PDF-1.4 content", content_type="application/pdf")
        self.client.post("/api/resume/", {"file": pdf}, format="multipart")
        active_resume = Resume.objects.filter(user=self.user).first()

        # 2. Setup connected email account and application
        EmailAccount.objects.create(
            user=self.user,
            email_address="sender@example.com",
            provider="gmail",
            access_token="token",
            refresh_token="refresh",
        )
        app = JobApplication.objects.create(
            user=self.user,
            company_name="Stripe",
            role_title="API Dev",
            recruiter_email="hr@stripe.com",
            status="sent",
        )

        # 3. Send email
        with patch("services.gmail_service.send_email") as mock_send:
            mock_send.return_value = "thread_resume_1"

            res = self.client.post(
                f"/api/job-applications/{app.id}/send-email/",
                {"subject": "Application", "body": "Please see attached resume."},
                format="json",
            )
            self.assertEqual(res.status_code, status.HTTP_200_OK)
            self.assertEqual(res.data["resume_attached"], "attached_resume.pdf")

            # Verify send_email was called with attachment_data
            mock_send.assert_called_once()
            call_kwargs = mock_send.call_args.kwargs
            self.assertEqual(call_kwargs["attachment_filename"], "attached_resume.pdf")
            self.assertEqual(call_kwargs["attachment_data"], b"%PDF-1.4 content")

            # Verify EmailLog records resume_attached
            log = EmailLog.objects.get(job_application=app)
            self.assertEqual(log.resume_attached, active_resume)




